"""
Export Utilities - Real Export Implementation with Proper Libraries

This module provides comprehensive export functionality using:
- openpyxl for Excel files
- python-docx for Word documents
- reportlab for PDF files
- csv for CSV files
"""

import csv
import io
from datetime import datetime
from typing import Dict, List, Any, Tuple
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Import our role-specific data generator
from utils.role_data import get_role_data

def format_number(value: Any) -> str:
    """Format numbers with thousand separators"""
    if isinstance(value, (int, float)):
        return f"{value:,}".replace(",", " ")
    return str(value)

def export_to_csv(export_type: str, role: str = None) -> io.BytesIO:
    """Export data to CSV format"""
    output = io.BytesIO()
    wrapper = io.TextIOWrapper(output, encoding='utf-8-sig', newline='')
    writer = csv.writer(wrapper)
    
    # Get role-specific data
    if not role:
        raise ValueError("Role parameter is required for export")
    
    data, headers = get_role_data(role, export_type)
    
    # Determine title based on role and export type
    titles = {
        "manager": {
            "orders": "Manager - Buyurtmalar hisoboti",
            "statistics": "Manager - Statistika hisoboti",
            "users": "Manager - Xodimlar hisoboti",
            "reports": "Manager - Hisobotlar"
        },
        "controller": {
            "orders": "Controller - Buyurtmalar hisoboti",
            "quality": "Controller - Sifat nazorati hisoboti",
            "users": "Controller - Texniklar hisoboti",
            "statistics": "Controller - Statistika hisoboti"
        },
        "call_center_supervisor": {
            "orders": "Call Center Supervisor - Buyurtmalar hisoboti",
            "users": "Call Center Supervisor - Xodimlar hisoboti",
            "feedback": "Call Center Supervisor - Fikr-mulohazalar",
            "workflow": "Call Center Supervisor - Workflow hisoboti",
            "statistics": "Call Center Supervisor - Statistika"
        },
        "admin": {
            "users": "Admin - Foydalanuvchilar hisoboti",
            "orders": "Admin - Buyurtmalar hisoboti",
            "system": "Admin - Tizim sozlamalari",
            "logs": "Admin - Tizim loglari",
            "statistics": "Admin - Umumiy statistika"
        },
        "warehouse": {
            "inventory": "Ombor - Inventarizatsiya hisoboti",
            "issued_items": "Ombor - Berilgan materiallar",
            "orders": "Ombor - Buyurtmalar hisoboti",
            "statistics": "Ombor - Statistika hisoboti"
        }
    }
    
    title = titles.get(role, {}).get(export_type, f"{role.title()} - {export_type.title()} hisoboti")
    
    # Write title and date
    writer.writerow([title])
    writer.writerow([f"Sana: {datetime.now().strftime('%Y-%m-%d %H:%M')}"])
    writer.writerow([])  # Empty row
    
    # Write headers
    writer.writerow(headers)
    
    # Write data based on type
    if isinstance(data, list):
        # List data (orders, users, etc.)
        for item in data:
            row = []
            for header in headers:
                # Map header to data key
                value = item.get(list(item.keys())[headers.index(header)]) if headers.index(header) < len(item) else "-"
                if isinstance(value, (int, float)) and not isinstance(value, bool):
                    value = format_number(value)
                row.append(value)
            writer.writerow(row)
    elif isinstance(data, dict):
        # Dictionary data (statistics, reports, etc.)
        if "overall_metrics" in data or "system_info" in data or any(isinstance(v, dict) for v in data.values()):
            # Complex nested structure
            for section, section_data in data.items():
                writer.writerow([])  # Empty row
                writer.writerow([section.replace("_", " ").title()])
                
                if isinstance(section_data, dict):
                    for key, value in section_data.items():
                        writer.writerow([key.replace("_", " ").title(), format_number(value) if isinstance(value, (int, float)) else value])
                elif isinstance(section_data, list):
                    for item in section_data:
                        if isinstance(item, dict):
                            writer.writerow([str(v) for v in item.values()])
                        else:
                            writer.writerow([item])
        else:
            # Simple key-value pairs
            for key, value in data.items():
                if isinstance(value, dict):
                    writer.writerow([key.replace("_", " ").title()])
                    for k, v in value.items():
                        writer.writerow(["", k.replace("_", " ").title(), format_number(v) if isinstance(v, (int, float)) else v])
                else:
                    writer.writerow([key.replace("_", " ").title(), format_number(value) if isinstance(value, (int, float)) else value])
    
    wrapper.flush()
    wrapper.detach()
    output.seek(0)
    return output

def export_to_excel(export_type: str, role: str = None) -> io.BytesIO:
    """Export data to Excel format with formatting"""
    output = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    
    # Get role-specific data
    if not role:
        raise ValueError("Role parameter is required for export")
    
    data, headers = get_role_data(role, export_type)
    
    # Set title
    titles = {
        "manager": {
            "orders": "Manager - Buyurtmalar",
            "statistics": "Manager - Statistika",
            "users": "Manager - Xodimlar",
            "reports": "Manager - Hisobotlar"
        },
        "controller": {
            "orders": "Controller - Buyurtmalar",
            "quality": "Controller - Sifat nazorati",
            "users": "Controller - Texniklar",
            "statistics": "Controller - Statistika"
        },
        "call_center_supervisor": {
            "orders": "Call Center - Buyurtmalar",
            "users": "Call Center - Xodimlar",
            "feedback": "Call Center - Fikr-mulohazalar",
            "workflow": "Call Center - Workflow",
            "statistics": "Call Center - Statistika"
        },
        "admin": {
            "users": "Admin - Foydalanuvchilar",
            "orders": "Admin - Buyurtmalar",
            "system": "Admin - Tizim",
            "logs": "Admin - Loglar",
            "statistics": "Admin - Statistika"
        },
        "warehouse": {
            "inventory": "Ombor - Inventarizatsiya",
            "issued_items": "Ombor - Berilgan materiallar",
            "orders": "Ombor - Buyurtmalar",
            "statistics": "Ombor - Statistika"
        }
    }
    
    ws.title = titles.get(role, {}).get(export_type, export_type.title())
    
    # Define styles
    header_font = Font(bold=True, color="FFFFFF", size=12)
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")
    
    # Add title and date
    ws.merge_cells('A1:' + get_column_letter(len(headers)) + '1')
    title_cell = ws['A1']
    title_cell.value = ws.title + " hisoboti"
    title_cell.font = Font(bold=True, size=16)
    title_cell.alignment = Alignment(horizontal="center")
    
    ws.merge_cells('A2:' + get_column_letter(len(headers)) + '2')
    date_cell = ws['A2']
    date_cell.value = f"Sana: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    date_cell.alignment = Alignment(horizontal="center")
    
    # Add headers
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col)
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
    
    # Add data
    row_num = 5
    if isinstance(data, list):
        # List data
        for item in data:
            for col, header in enumerate(headers, 1):
                value = list(item.values())[col-1] if col-1 < len(item) else "-"
                cell = ws.cell(row=row_num, column=col)
                if isinstance(value, (int, float)) and not isinstance(value, bool):
                    cell.value = value
                    cell.number_format = '#,##0'
                else:
                    cell.value = str(value)
            row_num += 1
    elif isinstance(data, dict):
        # Dictionary data
        for section, section_data in data.items():
            # Section header
            ws.merge_cells(f'A{row_num}:' + get_column_letter(len(headers)) + str(row_num))
            section_cell = ws.cell(row=row_num, column=1)
            section_cell.value = section.replace("_", " ").title()
            section_cell.font = Font(bold=True, size=14)
            section_cell.fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
            row_num += 1
            
            if isinstance(section_data, dict):
                for key, value in section_data.items():
                    ws.cell(row=row_num, column=1, value=key.replace("_", " ").title())
                    ws.cell(row=row_num, column=2, value=format_number(value) if isinstance(value, (int, float)) else str(value))
                    row_num += 1
            elif isinstance(section_data, list):
                for item in section_data:
                    if isinstance(item, dict):
                        for col, value in enumerate(item.values(), 1):
                            cell = ws.cell(row=row_num, column=col)
                            if isinstance(value, (int, float)) and not isinstance(value, bool):
                                cell.value = value
                                cell.number_format = '#,##0'
                            else:
                                cell.value = str(value)
                    row_num += 1
            row_num += 1  # Empty row between sections
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = get_column_letter(column[0].column)
        for cell in column:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Add borders
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    for row in ws.iter_rows(min_row=4):
        for cell in row:
            if cell.value:
                cell.border = thin_border
    
    wb.save(output)
    output.seek(0)
    return output

def export_to_word(export_type: str, role: str = None) -> io.BytesIO:
    """Export data to Word format"""
    doc = Document()
    
    # Get role-specific data
    if not role:
        raise ValueError("Role parameter is required for export")
    
    data, headers = get_role_data(role, export_type)
    
    # Set title
    titles = {
        "manager": {
            "orders": "Manager - Buyurtmalar hisoboti",
            "statistics": "Manager - Statistika hisoboti",
            "users": "Manager - Xodimlar hisoboti",
            "reports": "Manager - Hisobotlar"
        },
        "controller": {
            "orders": "Controller - Buyurtmalar hisoboti",
            "quality": "Controller - Sifat nazorati hisoboti",
            "users": "Controller - Texniklar hisoboti",
            "statistics": "Controller - Statistika hisoboti"
        },
        "call_center_supervisor": {
            "orders": "Call Center Supervisor - Buyurtmalar hisoboti",
            "users": "Call Center Supervisor - Xodimlar hisoboti",
            "feedback": "Call Center Supervisor - Fikr-mulohazalar",
            "workflow": "Call Center Supervisor - Workflow hisoboti",
            "statistics": "Call Center Supervisor - Statistika"
        },
        "admin": {
            "users": "Admin - Foydalanuvchilar hisoboti",
            "orders": "Admin - Buyurtmalar hisoboti",
            "system": "Admin - Tizim sozlamalari",
            "logs": "Admin - Tizim loglari",
            "statistics": "Admin - Umumiy statistika"
        },
        "warehouse": {
            "inventory": "Ombor - Inventarizatsiya hisoboti",
            "issued_items": "Ombor - Berilgan materiallar",
            "orders": "Ombor - Buyurtmalar hisoboti",
            "statistics": "Ombor - Statistika hisoboti"
        }
    }
    
    title = titles.get(role, {}).get(export_type, f"{role.title()} - {export_type.title()} hisoboti")
    
    # Add title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_para = doc.add_paragraph(f"Sana: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(12)
    
    doc.add_paragraph()  # Empty line
    
    # Handle different data types
    if isinstance(data, list) and data:
        # Create table for list data
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Add headers
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add data rows
        for item in data:
            row = table.add_row()
            for i, header in enumerate(headers):
                value = list(item.values())[i] if i < len(item) else "-"
                if isinstance(value, (int, float)) and not isinstance(value, bool):
                    value = format_number(value)
                row.cells[i].text = str(value)
        
        # Auto-fit table
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(1.2)
    
    elif isinstance(data, dict):
        # Handle dictionary data
        for section, section_data in data.items():
            # Add section heading
            doc.add_heading(section.replace("_", " ").title(), level=1)
            
            if isinstance(section_data, dict):
                # Create a simple table for key-value pairs
                table = doc.add_table(rows=0, cols=2)
                table.style = 'Light List Accent 1'
                
                for key, value in section_data.items():
                    row = table.add_row()
                    row.cells[0].text = key.replace("_", " ").title()
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
                    row.cells[1].text = format_number(value) if isinstance(value, (int, float)) else str(value)
                
            elif isinstance(section_data, list) and section_data:
                if isinstance(section_data[0], dict):
                    # Create table for list of dicts
                    headers = list(section_data[0].keys())
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # Add headers
                    for i, header in enumerate(headers):
                        cell = table.rows[0].cells[i]
                        cell.text = header.replace("_", " ").title()
                        cell.paragraphs[0].runs[0].font.bold = True
                    
                    # Add data
                    for item in section_data:
                        row = table.add_row()
                        for i, value in enumerate(item.values()):
                            if isinstance(value, (int, float)) and not isinstance(value, bool):
                                value = format_number(value)
                            row.cells[i].text = str(value)
            
            doc.add_paragraph()  # Empty line between sections
    
    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def export_to_pdf(export_type: str, role: str = None) -> io.BytesIO:
    """Export data to PDF format"""
    output = io.BytesIO()
    
    # Get role-specific data
    if not role:
        raise ValueError("Role parameter is required for export")
    
    data, headers = get_role_data(role, export_type)
    
    # Set title
    titles = {
        "manager": {
            "orders": "Manager - Buyurtmalar hisoboti",
            "statistics": "Manager - Statistika hisoboti",
            "users": "Manager - Xodimlar hisoboti",
            "reports": "Manager - Hisobotlar"
        },
        "controller": {
            "orders": "Controller - Buyurtmalar hisoboti",
            "quality": "Controller - Sifat nazorati hisoboti",
            "users": "Controller - Texniklar hisoboti",
            "statistics": "Controller - Statistika hisoboti"
        },
        "call_center_supervisor": {
            "orders": "Call Center Supervisor - Buyurtmalar hisoboti",
            "users": "Call Center Supervisor - Xodimlar hisoboti",
            "feedback": "Call Center Supervisor - Fikr-mulohazalar",
            "workflow": "Call Center Supervisor - Workflow hisoboti",
            "statistics": "Call Center Supervisor - Statistika"
        },
        "admin": {
            "users": "Admin - Foydalanuvchilar hisoboti",
            "orders": "Admin - Buyurtmalar hisoboti",
            "system": "Admin - Tizim sozlamalari",
            "logs": "Admin - Tizim loglari",
            "statistics": "Admin - Umumiy statistika"
        },
        "warehouse": {
            "inventory": "Ombor - Inventarizatsiya hisoboti",
            "issued_items": "Ombor - Berilgan materiallar",
            "orders": "Ombor - Buyurtmalar hisoboti",
            "statistics": "Ombor - Statistika hisoboti"
        }
    }
    
    title = titles.get(role, {}).get(export_type, f"{role.title()} - {export_type.title()} hisoboti")
    
    # Create PDF document
    if len(headers) > 8:
        doc = SimpleDocTemplate(output, pagesize=landscape(A4))
    else:
        doc = SimpleDocTemplate(output, pagesize=A4)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1f4788'),
        spaceAfter=30,
        alignment=1  # Center
    )
    
    # Add title
    elements.append(Paragraph(title, title_style))
    elements.append(Paragraph(f"Sana: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Handle different data types
    if isinstance(data, list) and data:
        # Prepare table data
        table_data = [headers]
        
        for item in data:
            row = []
            for i in range(len(headers)):
                value = list(item.values())[i] if i < len(item) else "-"
                if isinstance(value, (int, float)) and not isinstance(value, bool):
                    value = format_number(value)
                row.append(str(value))
            table_data.append(row)
        
        # Create table
        table = Table(table_data)
        
        # Style the table
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        
        elements.append(table)
    
    elif isinstance(data, dict):
        # Handle dictionary data
        for section, section_data in data.items():
            # Add section heading
            elements.append(Paragraph(section.replace("_", " ").title(), styles['Heading2']))
            elements.append(Spacer(1, 10))
            
            if isinstance(section_data, dict):
                # Create table for key-value pairs
                table_data = []
                for key, value in section_data.items():
                    table_data.append([
                        key.replace("_", " ").title(),
                        format_number(value) if isinstance(value, (int, float)) else str(value)
                    ])
                
                if table_data:
                    table = Table(table_data, colWidths=[3*inch, 3*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ]))
                    elements.append(table)
            
            elif isinstance(section_data, list) and section_data:
                if isinstance(section_data[0], dict):
                    # Create table for list of dicts
                    headers = list(section_data[0].keys())
                    table_data = [[h.replace("_", " ").title() for h in headers]]
                    
                    for item in section_data:
                        row = []
                        for value in item.values():
                            if isinstance(value, (int, float)) and not isinstance(value, bool):
                                value = format_number(value)
                            row.append(str(value))
                        table_data.append(row)
                    
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 11),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ]))
                    elements.append(table)
            
            elements.append(Spacer(1, 20))
    
    # Build PDF
    doc.build(elements)
    output.seek(0)
    return output

def create_export_file(export_type: str, format_type: str, role: str = None) -> Tuple[io.BytesIO, str]:
    """Create export file with real data using appropriate libraries"""
    if not role:
        raise ValueError("Role parameter is required for export")
    
    filename = get_export_filename(export_type, format_type, role)
    
    if format_type == "csv":
        file_content = export_to_csv(export_type, role)
    elif format_type == "xlsx":
        file_content = export_to_excel(export_type, role)
    elif format_type == "docx":
        file_content = export_to_word(export_type, role)
    elif format_type == "pdf":
        file_content = export_to_pdf(export_type, role)
    else:
        # Default to CSV
        file_content = export_to_csv(export_type, role)
        filename = filename.replace(f".{format_type}", ".csv")
    
    return file_content, filename

def get_export_filename(export_type: str, format_type: str, role: str = None) -> str:
    """Generate filename with timestamp"""
    if not role:
        raise ValueError("Role parameter is required for export")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{role}_{export_type}_{timestamp}.{format_type}"

def get_available_export_types(user_role: str) -> List[str]:
    """Get available export types based on user role"""
    if user_role == "warehouse":
        return ["inventory", "issued_items", "orders", "statistics"]
    elif user_role == "admin":
        return ["users", "orders", "statistics", "system", "logs"]
    elif user_role == "manager":
        return ["orders", "statistics", "users", "reports"]
    elif user_role == "controller":
        return ["orders", "statistics", "users", "quality", "reports"]
    elif user_role == "call_center_supervisor":
        return ["orders", "statistics", "users", "feedback", "workflow"]
    else:
        return ["orders", "statistics"]

def get_available_export_formats() -> List[str]:
    """Get available export formats"""
    return ["csv", "xlsx", "docx", "pdf"]